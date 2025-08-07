import logging
import asyncio
from pathlib import Path
from typing import List, Dict, Any, Optional, Union
import mimetypes
import hashlib

# Document processing imports
import PyPDF2
import fitz  # PyMuPDF
from docx import Document as DocxDocument
from PIL import Image

# LangChain imports
from langchain_core.documents import Document
from langchain_text_splitters import RecursiveCharacterTextSplitter

# Hebrew text processing
try:
    from hebrew_tokenizer import Tokenizer
    HEBREW_TOKENIZER_AVAILABLE = True
except ImportError:
    HEBREW_TOKENIZER_AVAILABLE = False

from config import settings
from services.ocr_service import ocr_service

logger = logging.getLogger(__name__)

class HebrewTextSplitter:
    """Enhanced text splitter for Hebrew documents"""
    
    def __init__(self):
        self.chunk_size = settings.chunk_size
        self.chunk_overlap = settings.chunk_overlap
        self.hebrew_tokenizer = None
        
        if HEBREW_TOKENIZER_AVAILABLE:
            self.hebrew_tokenizer = Tokenizer()
        
        # Initialize LangChain text splitter
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=self.chunk_size,
            chunk_overlap=self.chunk_overlap,
            length_function=self._hebrew_aware_length,
            separators=[
                "\n\n",  # Paragraph breaks
                "\n",    # Line breaks
                ".",     # Sentences
                "!",     # Exclamations
                "?",     # Questions
                ";",     # Semicolons
                ",",     # Commas
                " ",     # Spaces
                ""       # Characters
            ]
        )
    
    def _hebrew_aware_length(self, text: str) -> int:
        """Calculate text length considering Hebrew characters"""
        if self.hebrew_tokenizer and any('\u0590' <= c <= '\u05FF' for c in text):
            # Use Hebrew tokenizer for Hebrew text
            tokens = self.hebrew_tokenizer.tokenize(text)
            return len(tokens)
        else:
            # Use standard length for non-Hebrew text
            return len(text)
    
    def split_text(self, text: str, metadata: Optional[Dict] = None) -> List[Document]:
        """Split text into chunks with Hebrew awareness"""
        # Create a document from the text
        doc = Document(page_content=text, metadata=metadata or {})
        
        # Split using LangChain splitter
        chunks = self.text_splitter.split_documents([doc])
        
        # Add chunk metadata
        for i, chunk in enumerate(chunks):
            chunk.metadata.update({
                "chunk_id": i,
                "total_chunks": len(chunks),
                "chunk_size": len(chunk.page_content)
            })
        
        return chunks

class PDFProcessor:
    """PDF document processor"""
    
    @staticmethod
    async def extract_text(file_path: Path) -> Dict[str, Any]:
        """Extract text from PDF file"""
        try:
            text_content = []
            metadata = {
                "file_type": "pdf",
                "file_path": str(file_path),
                "pages": 0
            }
            
            # Try PyMuPDF first (better for complex layouts)
            try:
                doc = fitz.open(file_path)
                metadata["pages"] = len(doc)
                
                for page_num in range(len(doc)):
                    page = doc[page_num]
                    text = page.get_text()
                    if text.strip():
                        text_content.append({
                            "page": page_num + 1,
                            "content": text
                        })
                doc.close()
                
            except Exception as e:
                logger.warning(f"PyMuPDF failed for {file_path}, trying PyPDF2: {e}")
                
                # Fallback to PyPDF2
                with open(file_path, 'rb') as file:
                    pdf_reader = PyPDF2.PdfReader(file)
                    metadata["pages"] = len(pdf_reader.pages)
                    
                    for page_num, page in enumerate(pdf_reader.pages):
                        text = page.extract_text()
                        if text.strip():
                            text_content.append({
                                "page": page_num + 1,
                                "content": text
                            })
            
            return {
                "success": True,
                "content": text_content,
                "metadata": metadata,
                "error": None
            }
            
        except Exception as e:
            logger.error(f"Failed to process PDF {file_path}: {e}")
            return {
                "success": False,
                "content": [],
                "metadata": metadata,
                "error": str(e)
            }

class DocxProcessor:
    """DOCX document processor"""
    
    @staticmethod
    async def extract_text(file_path: Path) -> Dict[str, Any]:
        """Extract text from DOCX file"""
        try:
            doc = DocxDocument(file_path)
            
            text_content = []
            metadata = {
                "file_type": "docx",
                "file_path": str(file_path),
                "paragraphs": 0
            }
            
            for para in doc.paragraphs:
                if para.text.strip():
                    text_content.append(para.text)
            
            metadata["paragraphs"] = len(text_content)
            
            return {
                "success": True,
                "content": [{"content": "\n".join(text_content)}],
                "metadata": metadata,
                "error": None
            }
            
        except Exception as e:
            logger.error(f"Failed to process DOCX {file_path}: {e}")
            return {
                "success": False,
                "content": [],
                "metadata": {"file_type": "docx", "file_path": str(file_path)},
                "error": str(e)
            }

class ImageProcessor:
    """Image processor using OCR with layout awareness"""
    
    @staticmethod
    async def extract_text(file_path: Path) -> Dict[str, Any]:
        """Extract text from image using OCR with layout information"""
        try:
            # Use OCR service to extract text with layout information
            ocr_result = await ocr_service.extract_text_from_image(file_path, task_type="full")
            
            metadata = {
                "file_type": "image",
                "file_path": str(file_path),
                "ocr_method": ocr_result.get("method", "unknown"),
                "has_layout_info": ocr_result.get("parsed") is not None
            }
            
            if ocr_result["success"]:
                # Extract both raw text and structured layout data
                raw_text = ocr_result["text"]
                layout_data = ocr_result.get("parsed")
                
                # Create enhanced content with layout information
                content_items = []
                
                if layout_data and isinstance(layout_data, dict):
                    # Process structured layout data
                    layout_elements = layout_data.get("elements", [])
                    
                    # Group elements by type for better processing
                    text_elements = []
                    table_elements = []
                    formula_elements = []
                    
                    for element in layout_elements:
                        element_type = element.get("category", "Text")
                        element_text = element.get("text", "")
                        
                        if element_type == "Table":
                            table_elements.append({
                                "type": "table",
                                "content": element_text,
                                "bbox": element.get("bbox"),
                                "category": element_type
                            })
                        elif element_type == "Formula":
                            formula_elements.append({
                                "type": "formula",
                                "content": element_text,
                                "bbox": element.get("bbox"),
                                "category": element_type
                            })
                        else:
                            text_elements.append({
                                "type": "text",
                                "content": element_text,
                                "bbox": element.get("bbox"),
                                "category": element_type
                            })
                    
                    # Add structured content
                    if text_elements:
                        content_items.append({
                            "type": "text_block",
                            "content": "\n".join([elem["content"] for elem in text_elements]),
                            "metadata": {
                                "element_count": len(text_elements),
                                "categories": list(set([elem["category"] for elem in text_elements]))
                            }
                        })
                    
                    if table_elements:
                        for table in table_elements:
                            content_items.append({
                                "type": "table",
                                "content": table["content"],
                                "metadata": {
                                    "category": table["category"],
                                    "bbox": table["bbox"]
                                }
                            })
                    
                    if formula_elements:
                        for formula in formula_elements:
                            content_items.append({
                                "type": "formula",
                                "content": formula["content"],
                                "metadata": {
                                    "category": formula["category"],
                                    "bbox": formula["bbox"]
                                }
                            })
                    
                    # Add layout metadata
                    metadata.update({
                        "layout_elements": len(layout_elements),
                        "text_elements": len(text_elements),
                        "table_elements": len(table_elements),
                        "formula_elements": len(formula_elements),
                        "layout_data": layout_data
                    })
                
                # Always include raw text as fallback
                if not content_items:
                    content_items.append({
                        "type": "raw_text",
                        "content": raw_text,
                        "metadata": {"source": "ocr_raw"}
                    })
                
                return {
                    "success": True,
                    "content": content_items,
                    "metadata": metadata,
                    "error": None
                }
            else:
                return {
                    "success": False,
                    "content": [],
                    "metadata": metadata,
                    "error": ocr_result.get("error", "OCR failed")
                }
                
        except Exception as e:
            logger.error(f"Failed to process image {file_path}: {e}")
            return {
                "success": False,
                "content": [],
                "metadata": {"file_type": "image", "file_path": str(file_path)},
                "error": str(e)
            }

class AudioProcessor:
    """Audio processor using transcription service"""
    
    @staticmethod
    async def extract_text(file_path: Path) -> Dict[str, Any]:
        """Extract text from audio file using transcription"""
        try:
            # Check if transcription service is available
            try:
                from services import transcription_service
                transcription_available = transcription_service is not None
            except ImportError:
                transcription_available = False
                logger.warning("Transcription service not available - audio processing disabled")
            
            if not transcription_available:
                return {
                    "success": False,
                    "content": [],
                    "metadata": {
                        "file_type": "audio", 
                        "file_path": str(file_path),
                        "error": "Transcription service not available"
                    },
                    "error": "Transcription service not available"
                }
            
            # Initialize transcription service if needed
            if not transcription_service.initialized:
                try:
                    await transcription_service.initialize()
                except Exception as e:
                    logger.error(f"Failed to initialize transcription service: {e}")
                    return {
                        "success": False,
                        "content": [],
                        "metadata": {
                            "file_type": "audio", 
                            "file_path": str(file_path),
                            "error": "Transcription service initialization failed"
                        },
                        "error": f"Transcription service initialization failed: {e}"
                    }
            
            # Transcribe audio file
            transcription_result = await transcription_service.transcribe_audio_file(
                str(file_path)
            )
            
            metadata = {
                "file_type": "audio",
                "file_path": str(file_path),
                "transcription_method": "whisper-hebrew"
            }
            
            if transcription_result["success"]:
                return {
                    "success": True,
                    "content": [{"content": transcription_result["text"]}],
                    "metadata": metadata,
                    "error": None
                }
            else:
                return {
                    "success": False,
                    "content": [],
                    "metadata": metadata,
                    "error": transcription_result.get("error", "Transcription failed")
                }
                
        except Exception as e:
            logger.error(f"Failed to process audio {file_path}: {e}")
            return {
                "success": False,
                "content": [],
                "metadata": {"file_type": "audio", "file_path": str(file_path)},
                "error": str(e)
            }

class DocumentProcessor:
    """Main document processing service"""
    
    def __init__(self):
        self.text_splitter = HebrewTextSplitter()
        self.processors = {
            "pdf": PDFProcessor,
            "docx": DocxProcessor,
            "doc": DocxProcessor,  # Treat DOC as DOCX
            "txt": self._process_text_file,
            "image": ImageProcessor,
            "audio": AudioProcessor
        }
        
    async def initialize(self):
        """Initialize document processor"""
        # Initialize OCR service if not already done
        if not ocr_service.initialized:
            await ocr_service.initialize()
        logger.info("Document processor initialized")
    
    def _get_file_type(self, file_path: Path) -> str:
        """Determine file type from extension"""
        extension = file_path.suffix.lower()
        
        type_mapping = {
            '.pdf': 'pdf',
            '.docx': 'docx',
            '.doc': 'doc',
            '.txt': 'txt',
            '.png': 'image',
            '.jpg': 'image',
            '.jpeg': 'image',
            '.tiff': 'image',
            '.bmp': 'image',
            '.mp3': 'audio',
            '.wav': 'audio',
            '.m4a': 'audio',
            '.flac': 'audio'
        }
        
        return type_mapping.get(extension, 'unknown')
    
    async def _process_text_file(self, file_path: Path) -> Dict[str, Any]:
        """Process plain text file"""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                content = file.read()
            
            return {
                "success": True,
                "content": [{"content": content}],
                "metadata": {
                    "file_type": "txt",
                    "file_path": str(file_path)
                },
                "error": None
            }
            
        except Exception as e:
            logger.error(f"Failed to process text file {file_path}: {e}")
            return {
                "success": False,
                "content": [],
                "metadata": {"file_type": "txt", "file_path": str(file_path)},
                "error": str(e)
            }
    
    async def process_document(self, file_path: Union[str, Path]) -> Dict[str, Any]:
        """Process a single document"""
        file_path = Path(file_path)
        
        if not file_path.exists():
            return {
                "success": False,
                "error": f"File not found: {file_path}",
                "documents": [],
                "metadata": {}
            }
        
        # Check file size
        file_size_mb = file_path.stat().st_size / (1024 * 1024)
        if file_size_mb > settings.max_file_size_mb:
            return {
                "success": False,
                "error": f"File too large: {file_size_mb:.1f}MB (max: {settings.max_file_size_mb}MB)",
                "documents": [],
                "metadata": {}
            }
        
        # Determine file type and processor
        file_type = self._get_file_type(file_path)
        
        if file_type == 'unknown':
            return {
                "success": False,
                "error": f"Unsupported file type: {file_path.suffix}",
                "documents": [],
                "metadata": {}
            }
        
        # Process the document
        try:
            if file_type == 'txt':
                extraction_result = await self._process_text_file(file_path)
            else:
                processor = self.processors[file_type]
                extraction_result = await processor.extract_text(file_path)
            
            if not extraction_result["success"]:
                return {
                    "success": False,
                    "error": extraction_result["error"],
                    "documents": [],
                    "metadata": extraction_result["metadata"]
                }
            
            # Combine all extracted text with layout awareness
            full_text = ""
            layout_metadata = {}
            
            for content_item in extraction_result["content"]:
                if isinstance(content_item, dict):
                    content_type = content_item.get("type", "text")
                    content = content_item.get("content", "")
                    
                    # Add type-specific formatting
                    if content_type == "table":
                        full_text += f"\n[טבלה]\n{content}\n[/טבלה]\n\n"
                        layout_metadata["has_tables"] = True
                    elif content_type == "formula":
                        full_text += f"\n[נוסחה]\n{content}\n[/נוסחה]\n\n"
                        layout_metadata["has_formulas"] = True
                    elif content_type == "text_block":
                        full_text += f"{content}\n\n"
                        layout_metadata["has_text_blocks"] = True
                    else:
                        full_text += f"{content}\n\n"
                else:
                    full_text += str(content_item) + "\n\n"
            
            # Add document hash for deduplication
            doc_hash = hashlib.md5(full_text.encode()).hexdigest()
            
            # Create base metadata with layout information
            base_metadata = extraction_result["metadata"].copy()
            base_metadata.update({
                "document_hash": doc_hash,
                "file_size_mb": file_size_mb,
                "total_chars": len(full_text),
                "layout_aware": True
            })
            
            # Add layout-specific metadata
            if layout_metadata:
                base_metadata.update(layout_metadata)
            
            # Add OCR-specific metadata if available
            if extraction_result["metadata"].get("has_layout_info"):
                base_metadata.update({
                    "ocr_enhanced": True,
                    "layout_elements": extraction_result["metadata"].get("layout_elements", 0),
                    "text_elements": extraction_result["metadata"].get("text_elements", 0),
                    "table_elements": extraction_result["metadata"].get("table_elements", 0),
                    "formula_elements": extraction_result["metadata"].get("formula_elements", 0)
                })
            
            # Split text into chunks
            chunks = self.text_splitter.split_text(full_text, base_metadata)
            
            return {
                "success": True,
                "error": None,
                "documents": chunks,
                "metadata": base_metadata
            }
            
        except Exception as e:
            logger.error(f"Failed to process document {file_path}: {e}")
            return {
                "success": False,
                "error": str(e),
                "documents": [],
                "metadata": {"file_path": str(file_path)}
            }
    
    async def process_multiple_documents(self, file_paths: List[Union[str, Path]]) -> Dict[str, Any]:
        """Process multiple documents concurrently with memory management"""
        results = {
            "successful": [],
            "failed": [],
            "total_documents": 0,
            "total_chunks": 0,
            "memory_usage": {}
        }
        
        # Process documents in batches to manage memory
        batch_size = 5  # Process 5 documents at a time
        total_files = len(file_paths)
        
        logger.info(f"Processing {total_files} documents in batches of {batch_size}")
        
        for i in range(0, total_files, batch_size):
            batch_files = file_paths[i:i + batch_size]
            batch_num = (i // batch_size) + 1
            total_batches = (total_files + batch_size - 1) // batch_size
            
            logger.info(f"Processing batch {batch_num}/{total_batches} ({len(batch_files)} files)")
            
            # Process current batch
            tasks = [self.process_document(path) for path in batch_files]
            batch_results = await asyncio.gather(*tasks, return_exceptions=True)
            
            # Process batch results
            for file_path, result in zip(batch_files, batch_results):
                if isinstance(result, Exception):
                    results["failed"].append({
                        "file_path": str(file_path),
                        "error": str(result)
                    })
                elif result["success"]:
                    results["successful"].append({
                        "file_path": str(file_path),
                        "documents": result["documents"],
                        "metadata": result["metadata"]
                    })
                    results["total_chunks"] += len(result["documents"])
                else:
                    results["failed"].append({
                        "file_path": str(file_path),
                        "error": result["error"]
                    })
            
            # Memory cleanup after each batch
            await self._cleanup_batch_memory()
            
            # Log progress
            processed_so_far = len(results["successful"]) + len(results["failed"])
            logger.info(f"Progress: {processed_so_far}/{total_files} documents processed")
        
        results["total_documents"] = len(results["successful"])
        
        # Final memory usage report
        results["memory_usage"] = await self._get_memory_usage()
        
        return results
    
    async def _cleanup_batch_memory(self):
        """Clean up memory after processing a batch"""
        try:
            # Clean up OCR service memory if available
            if hasattr(ocr_service, 'dots_ocr') and ocr_service.dots_ocr.is_available():
                await ocr_service.dots_ocr.cleanup_memory()
            
            # Force garbage collection
            import gc
            gc.collect()
            
            logger.debug("Batch memory cleanup completed")
            
        except Exception as e:
            logger.warning(f"Error during batch memory cleanup: {e}")
    
    async def _get_memory_usage(self) -> Dict[str, Any]:
        """Get current memory usage information"""
        try:
            memory_info = {}
            
            # Get OCR service memory usage
            if hasattr(ocr_service, 'dots_ocr') and ocr_service.dots_ocr.is_available():
                memory_info["ocr_service"] = ocr_service.dots_ocr.get_memory_usage()
            
            # Get system memory usage
            import psutil
            memory_info["system"] = {
                "memory_percent": psutil.virtual_memory().percent,
                "memory_available": f"{psutil.virtual_memory().available / 1024**3:.2f} GB",
                "memory_total": f"{psutil.virtual_memory().total / 1024**3:.2f} GB"
            }
            
            return memory_info
            
        except Exception as e:
            logger.warning(f"Error getting memory usage: {e}")
            return {"error": str(e)}

# Global document processor instance
document_processor = DocumentProcessor()